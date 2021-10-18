import PIL
import docx
import glob
import matplotlib.pyplot as plt
import numpy as np
import os
import pandas as pd
import pathlib
import platform
import re
import requests
import shutil
import subprocess
import sys
import wx
import wx.adv
import wx.lib.scrolledpanel as scrolled
# from IPython.core.display import HTML
from PIL import Image
from bs4 import BeautifulSoup
from datetime import date
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Mm, Cm, Pt
from io import StringIO, BytesIO
from pkg_resources import resource_filename
from threading import Thread
import matplotlib
matplotlib.use('Agg')


def mainGui():
    app = wx.App()

    try:
        splashbmp= os.path.abspath(resource_filename('analizatorkotow', 'splash.png'))
        bitmap = wx.Bitmap(splashbmp)
        splash = wx.adv.SplashScreen(bitmap, wx.adv.SPLASH_CENTRE_ON_SCREEN | wx.adv.SPLASH_TIMEOUT, 4500, None, -1)

    except:
        bitmap = wx.Bitmap('analizatorkotow/splash.png')
        splash = wx.adv.SplashScreen(bitmap, wx.adv.SPLASH_CENTRE_ON_SCREEN | wx.adv.SPLASH_TIMEOUT, 4500, None, -1)

    splash.Show()
    if os.name == 'nt':
        app.locale = wx.Locale(wx.LANGUAGE_POLISH)

    CatFrame().Show()
    app.MainLoop()


def boldframe(s):

    if 'Ochota' in s.Miejsce:
        return ['font-weight: bold'] * len(s.index)
    else:
        return ['font-weight: normal'] * len(s.index)

def Tooltipset(object, tooltip):
    if float(wx.__version__[0]) < 4:
        object.SetToolTipString(tooltip)
    else:
        object.SetToolTip(tooltip)


class GeneralLayout(scrolled.ScrolledPanel):
    def __init__(self, parent):



        self.newgen=[False, False]


        self.my_datadir = self.get_datadir() / "kotolizator" #generates path for application support to store backups
        self.today = date.today().strftime("%d-%m-%Y")
        print(self.today)
        try:
            self.my_datadir.mkdir(parents=True)
        except FileExistsError:
            pass
        print(self.my_datadir)
        if sys.platform == "win32":
            slash = "\\"
        else:
            slash='/'

        self.my_datadir2=str(self.my_datadir)+slash
        self.nadzis='Stan na dziś'
        # self.listbackups  = glob.glob(self.my_datadir2 + 'Mia*.kotolizer').split('Miau')[1].split('.kotolizer')[0]
        # print(folder)
        scrolled.ScrolledPanel.__init__(self, parent)
        self.SetupScrolling(scroll_y=True, scroll_x=False)
        self.generalboxver = wx.BoxSizer(wx.VERTICAL)  # General vertical box for UI
        self.word1 = wx.Button(self, label='Wygeneruj zestawienie', size=(200, 50))
        self.word1.Enabled = False

        self.box1 = wx.BoxSizer(wx.HORIZONTAL)
        self.box2 = wx.BoxSizer(wx.HORIZONTAL)
        self.box2a = wx.BoxSizer(wx.HORIZONTAL)
        self.box2b = wx.BoxSizer(wx.HORIZONTAL)
        self.box4 = wx.BoxSizer(wx.HORIZONTAL)

        self.box3 = wx.BoxSizer(wx.HORIZONTAL)
        self.box3a = wx.BoxSizer(wx.VERTICAL)
        self.box5 = wx.BoxSizer(wx.HORIZONTAL)
        self.box6 = wx.BoxSizer(wx.HORIZONTAL)
        self.generalboxver.Add(self.box1, 0, wx.ALL | wx.EXPAND, 1)

        self.generalboxver.Add(self.box2, 0, wx.ALL | wx.EXPAND, 1)

        self.generalboxver.Add(self.box2a,0,wx.ALL | wx.EXPAND, 1)
        self.generalboxver.Add(self.box2b,0,wx.ALL | wx.EXPAND, 1)

        self.generalboxver.Add(self.box3,0,wx.ALL | wx.EXPAND, 1)
        self.generalboxver.Add(self.box3a,0,wx.ALL | wx.EXPAND, 1)
        self.generalboxver.Add(self.box4,0,wx.ALL | wx.EXPAND, 1)

        self.generalboxver.Add(self.box5 ,0,wx.ALL | wx.EXPAND, 1)
        self.generalboxver.Add(self.box6,0, wx.ALL | wx.EXPAND, 1)
        # self.generalboxver.Add(self.box2a, 0, wx.ALL | wx.EXPAND, 1)
        self.koty = wx.ToggleButton(self, label='Analizuj koty')
        self.psy = wx.ToggleButton(self, label='Analizuj psy')
        self.backup = wx.ComboBox(self, choices=[], style=wx.CB_READONLY)
        self.backup2 = wx.ComboBox(self, choices=[], style=wx.CB_READONLY)
        self.histogram1 = wx.Button(self, label='Wyświetl histogram', size=(200, 50))
        self.histogram1.Enabled = False
        self.genxls = wx.Button(self, label='Tabelka', size=(200, 50))
        self.genxls.Enabled = False
        self.statcat = wx.Button(self, label='Miejsca', size=(200, 50))
        self.statcat.Enabled = False
        # Definition of buttons
        self.kotybool = False
        self.histnew = wx.Button(self, label='Histogram nowych zwierząt', size=(200, 50))
        self.histold = wx.Button(self, label='Histogram znikniętych zwierząt', size=(200, 50))

        self.wordNew = wx.Button(self, label='Wygeneruj zestawienie nowych zwierząt', size=(200, 50))
        self.wordOld = wx.Button(self, label='Wygeneruj zestawienie znikniętych zwierząt', size=(200, 50))

        self.switchkotpies(True)

        # self.listbackups = ['Stan na di']
        # self.listbackups.insert(0,'Stan na dziś')
        # self.zwierzeta = ['Analizuj koty', 'Analizuj psy', 'Analizuj wszystkie']
        # self.zwierzetacombo = wx.ComboBox(self, choices=self.zwierzeta, style=wx.CB_READONLY)
        # Tooltipset(self.zwierzetacombo, 'Wybierz gatunek do analizy')
        # self.zwierzetacombo.Value = self.zwierzeta[0]
        self.box1.Add(self.koty, wx.ALL | wx.EXPAND, 1)
        self.box1.Add(self.psy, wx.ALL | wx.EXPAND, 1)

        Tooltipset(self.backup, 'Na jaki dzień wyświetlić dane?')
        # self.backup.Value = self.listbackups[0]
        self.koty.SetValue(True)
        self.psy.SetValue(False)
        self.koty.Bind(wx.EVT_TOGGLEBUTTON, self.switchkotpies)
        self.psy.Bind(wx.EVT_TOGGLEBUTTON, self.switchkotpies)

        self.nowedane = wx.Button(self, label='Ściągnij nowe dane', size=(200, 50))

        self.nowedane.Bind(wx.EVT_BUTTON, self.parsepaluch)
        self.word1.Bind(wx.EVT_BUTTON, self.genword)


        self.box2.Add(self.nowedane)
        self.box2.Add(self.backup, wx.ALL | wx.EXPAND, 1)

        self.box2a.Add(self.word1, wx.ALL | wx.EXPAND, 1)
        self.box2a.Add(self.histogram1, wx.ALL | wx.EXPAND, 1)
        self.box2b.Add(self.genxls, wx.ALL | wx.EXPAND, 1)
        self.box2b.Add(self.statcat, wx.ALL | wx.EXPAND, 1)

        self.postep = wx.Gauge(self, range=100, size=(600, 50), name='Postęp')
        Tooltipset(self.postep, 'Pasek postępu ściągania nowych danych, lub generacji zestawienia')
        self.box3.Add(self.postep, wx.ALL | wx.EXPAND, 1)
        # self.generalboxver.Add(wx.StaticLine(self, style=wx.LI_HORIZONTAL))

        # self.generalboxver.Add(wx.StaticText(self, label='Porównania historyczne: \n'))

        # self.generalboxver.Add(wx.StaticLine(self, style=wx.LI_HORIZONTAL))
        # self.dialog = wx.ProgressDialog('Doing Stuff', 'Please wait...')

        Tooltipset(self.backup2, 'Na jaki dzień wyświetlić dane?')
        self.backup2.Value = self.listbackups[0]
        # self.generalboxver.Add(wx.StaticLine(self, style=wx.LI_HORIZONTAL))

        self.txtpor = wx.StaticText(self, label='Porównaj z:')
        self.txtpor2 = wx.StaticText(self, label='*********************************************************Kącik historyczny:*********************************************************')

        self.box4.Add(self.txtpor)
        self.box3a.Add(wx.StaticText(self, label=' '))

        self.box3a.Add(self.txtpor2)
        self.box3a.Add(wx.StaticText(self, label=' '))

        self.box4.Add(self.backup2, wx.ALL | wx.EXPAND, 1)

        self.box5.Add(self.wordNew, wx.ALL | wx.EXPAND, 1)
        self.box6.Add(self.wordOld, wx.ALL | wx.EXPAND, 1)
        self.backup.Bind(wx.EVT_COMBOBOX, self.combochange)
        self.wordNew.Bind(wx.EVT_BUTTON, self.genwordNew)
        self.wordOld.Bind(wx.EVT_BUTTON, self.genwordOld)
        self.box5.Add(self.histnew, wx.ALL | wx.EXPAND, 1)
        self.box6.Add(self.histold, wx.ALL | wx.EXPAND, 1)

        self.histogram1.Bind(wx.EVT_BUTTON, self.plothist)
        self.genxls.Bind(wx.EVT_BUTTON, self.excellgen)
        self.statcat.Bind(wx.EVT_BUTTON, self.getcatstat)
        self.histnew.Bind(wx.EVT_BUTTON, self.plothistnew)
        self.histold.Bind(wx.EVT_BUTTON, self.plothistold)
        self.lab = True
        self.SetSizerAndFit(self.generalboxver)

    def getcatstat(self,event):
        xlscat=self.gencatxls()

        dlg = wx.MessageDialog(self,  xlscat.groupby(['Miejsce']).size().to_string(),'Statystyki',).ShowModal()
    def gencatxls(self):

        catframe = self.getframecur(self.backup)
        xlscat = pd.DataFrame(columns=['Nr', 'ID', 'Imię', 'Wiek', 'Miejsce'])
        opo=0
        for a, b in catframe.iterrows():
            opo+=1
            Imie = b[2].split('\n')[0]
            Miejsce = b[4].split('\n')[-1]
            miejscatocheck = ['Azyl', 'Ochota', 'właściciela']
            if not any(ext in Miejsce for ext in miejscatocheck):
                Miejsce = 'brak informacji'
            if 'mies' in b[3]:
                mies=int(b[3].split()[0])
                if mies<4:
                    Imie +='(m)'
            new_row = {'Nr':opo, 'ID': b[1], 'Imię': Imie, 'Wiek': b[3], 'Miejsce': Miejsce}
            xlscat = xlscat.append(new_row, ignore_index=True)
        return xlscat
    def excellgen(self,event):
        xlscat=self.gencatxls()
        # xlscat.style.applymap('font-weight: bold',
        #                        subset=pd.IndexSlice[xlscat.index[xlscat.Miejsce == 'Ochota na Kota'], :])
        # print(xlscat.style.applymap('font-weight: bold',
        #                        subset=pd.IndexSlice[xlscat.index[xlscat.Miejsce == 'Ochota na Kota'], :]))
        with wx.FileDialog(self, "Zapisz tabelkę", wildcard="`XLSX files (*.xlsx)|*.xlsx",
                           style=wx.FD_SAVE | wx.FD_OVERWRITE_PROMPT) as fileDialog:
            if fileDialog.ShowModal() == wx.ID_CANCEL:
                return  # the user changed their mind

            # save the current contents in the file
            pathname = fileDialog.GetPath()
            with pd.ExcelWriter(pathname) as writer:
                xlscat.style.apply(boldframe,axis=1).to_excel(writer, sheet_name='Koty', index=False)
            self.openfile(pathname)

    def gencathist(self,catframe, ax1, ax2, sex,kotbool):


        lataM = []
        lataF = []
        lata = []
        for k in zip(catframe['Wiek'], catframe['Imię']):
            if 'samiec' in k[1]:
                if 'rok' in k[0] or 'lat' in k[0]:
                    lataM.append(int(k[0][0:2]) * 12)
                else:
                    lataM.append(int(k[0][0:2]))
            elif 'samica' in k[1]:
                if 'rok' in k[0] or 'lat' in k[0]:
                    lataF.append(int(k[0][0:2]) * 12)
                else:
                    lataF.append(int(k[0][0:2]))

            if 'rok' in k[0] or 'lat' in k[0]:
                lata.append(int(k[0][0:2]) * 12)
            else:
                lata.append(int(k[0][0:2]))


        bins = np.arange(0, 13) - 0.5

        if sex == 'all':
            lataR = np.floor(np.array(lata) / 12)
            oldest = int(np.max(lataR))
            bins2 = np.arange(0, oldest + 2) - 0.5

            a1, _, patches = ax1.hist(np.array(lata), bins, linewidth=1, edgecolor='w', color='k')
            a2, _, patches2 = ax2.hist(lataR, bins2, linewidth=1, edgecolor='w', color='k')
            if kotbool:
                ax1.set_title('Koty do 1 roku, suma: ' + str(round(sum(a1))))
                ax2.set_title('Koty wszystkie, suma: ' + str(round(sum(a2))) + ';starsze of 8lat: ' + str(
                    len(np.where(lataR >= 8)[0])))
            else:
                ax1.set_title('Psy do 1 roku, suma: ' + str(round(sum(a1))))
                ax2.set_title('Psy wszystkie, suma: ' + str(round(sum(a2))) + ';starsze of 8lat: ' + str(
                    len(np.where(lataR >= 8)[0])))

        elif sex == 'f':
            lataR = np.floor(np.array(lataF) / 12)

            oldest = int(np.max(lataR))
            bins2 = np.arange(0, oldest + 2) - 0.5

            a1, _, patches = ax1.hist(np.array(lataF), bins, linewidth=1, edgecolor='w', color='pink')
            a2, _, patches2 = ax2.hist(lataR, bins2, linewidth=1, edgecolor='w', color='pink')

            if kotbool:

                ax1.set_title('Kotki do 1 roku, suma: ' + str(round(sum(a1))))
                ax2.set_title('Kotki wszystkie, suma: ' + str(round(sum(a2))) + ';starsze of 8lat: ' + str(
                    len(np.where(lataR >= 8)[0])))
            else:
                ax1.set_title('Suczki do 1 roku, suma: ' + str(round(sum(a1))))
                ax2.set_title('Suczki wszystkie, suma: ' + str(round(sum(a2))) + ';starsze of 8lat: ' + str(
                    len(np.where(lataR >= 8)[0])))


        else:
            lataR = np.floor(np.array(lataM) / 12)

            oldest = int(np.max(lataR))
            bins2 = np.arange(0, oldest + 2) - 0.5

            a1, _, patches = ax1.hist(np.array(lataM), bins, linewidth=1, edgecolor='w', color='b')
            # ax1.set_title('Kocury do 1 roku, suma: ' + str(round(sum(a1))))
            a2, _, patches2 = ax2.hist(lataR, bins2, linewidth=1, edgecolor='w', color='b')

            # ax2.set_title('Kocury wszystkie, suma: ' + str(round(sum(a2))) + ';starsze of 8lat:' + str(
            # len(np.where(lataR >= 8)[0])))
            if kotbool:

                ax1.set_title('Kocury do 1 roku, suma: ' + str(round(sum(a1))))
                ax2.set_title('Kocury wszystkie, suma: ' + str(round(sum(a2))) + ';starsze of 8lat: ' + str(
                    len(np.where(lataR >= 8)[0])))
            else:
                ax1.set_title('Psy (samce) do 1 roku, suma: ' + str(round(sum(a1))))
                ax2.set_title('Psy (samce) wszystkie, suma: ' + str(round(sum(a2))) + ';starsze of 8lat: ' + str(
                    len(np.where(lataR >= 8)[0])))

        ax1.set_xticks(range(0, 12))

        for pp in patches:
            x = (pp._x0 + pp._x1) / 2
            y = pp._y1 + 0.3
            ax1.text(x, y, int(pp._y1), ha='center')

        #     ax1.set_title('Koty do 1 roku, suma: '  + str(round(sum(a1))))
        ax1.set_xlabel('miesiące')
        lataR = np.floor(np.array(lata) / 12)

        #     oldest=int(np.max(lataR))
        #     bins = np.arange(0,oldest+2) -0.5

        #     a2, _, patches=ax2.hist(lataR,bins, linewidth=1, edgecolor='w')
        ax2.set_xticks(range(0, oldest + 1))

        ax2.set_xlabel('lata')
        if kotbool:
            ax2.set_ylabel('Liczba kotów')
            ax1.set_ylabel('Liczba kotów')
        else:
            ax2.set_ylabel('Liczba psów')
            ax1.set_ylabel('Liczba psów')

        for pp in patches2:
            x = (pp._x0 + pp._x1) / 2
            y = pp._y1 + 0.3
            ax2.text(x, y, int(pp._y1), ha='center')

        return lata

    def plothist(self,event):
        with wx.FileDialog(self, "Save histogram", wildcard="PNG files (*.png)|*.png",
                           style=wx.FD_SAVE | wx.FD_OVERWRITE_PROMPT) as fileDialog:
            if fileDialog.ShowModal() == wx.ID_CANCEL:
                return  # the user changed their mind

            # save the current contents in the file
            pathname = fileDialog.GetPath()
            catframe=self.getframecur(self.backup)

            first2 = Thread(target=self.plothistp, args=(pathname,catframe))
            first2.setDaemon(True)
            first2.start()
            print(pathname)
    def plothistold(self,event):
        catframe = self.getframecur(self.backup)
        catframe2 = self.getframecur(self.backup2)
        # Nowekoty = comparedf(catframe, zeszly_tydzien)
        Zniknietekoty = self.comparedf(catframe2, catframe)
        if Zniknietekoty.shape[0] == 0:
            dlg = wx.MessageDialog(self, 'Nie ubyło zwierzaków', 'Stan zwierząt nie zmniejszył się').ShowModal()
        else:
            with wx.FileDialog(self, "Save histogram", wildcard="PNG files (*.png)|*.png",
                               style=wx.FD_SAVE | wx.FD_OVERWRITE_PROMPT) as fileDialog:
                if fileDialog.ShowModal() == wx.ID_CANCEL:
                    return  # the user changed their mind

                # save the current contents in the file
                pathname = fileDialog.GetPath()

                first2 = Thread(target=self.plothistp, args=(pathname,Zniknietekoty))
                first2.setDaemon(True)
                first2.start()
                print(pathname)
    def plothistnew(self,event):
        catframe = self.getframecur(self.backup)
        catframe2 = self.getframecur(self.backup2)
        # Nowekoty = comparedf(catframe, zeszly_tydzien)
        Nowekoty = self.comparedf(catframe, catframe2)
        if Nowekoty.shape[0]==0:
            dlg = wx.MessageDialog(self, 'Nie przybyło zwierzaków', 'Nie ma nowych zwierząt w porówaniu').ShowModal()
        else:
            with wx.FileDialog(self, "Save histogram", wildcard="PNG files (*.png)|*.png",
                               style=wx.FD_SAVE | wx.FD_OVERWRITE_PROMPT) as fileDialog:
                if fileDialog.ShowModal() == wx.ID_CANCEL:
                    return  # the user changed their mind

                # save the current contents in the file
                pathname = fileDialog.GetPath()

                first2 = Thread(target=self.plothistp, args=(pathname, Nowekoty))
                first2.setDaemon(True)
                first2.start()
                print(pathname)


    def plothistp(self,pathname,catframe):
        # print(catframe.shape)
        fig = plt.figure(figsize=(29, 25))
        SMALL_SIZE = 14
        MEDIUM_SIZE = 25
        BIGGER_SIZE = 32
        ax1 = fig.add_subplot(321)
        ax2 = fig.add_subplot(322)
        ax3 = fig.add_subplot(323)
        ax4 = fig.add_subplot(324)
        ax5 = fig.add_subplot(325)
        ax6 = fig.add_subplot(326)
        wx.CallAfter(self.postep.SetRange, 3)

        plt.rc('font', size=SMALL_SIZE)  # controls default text sizes
        plt.rc('axes', titlesize=BIGGER_SIZE)  # fontsize of the axes title
        plt.rc('axes', labelsize=MEDIUM_SIZE)  # fontsize of the x and y labels
        plt.rc('xtick', labelsize=SMALL_SIZE)  # fontsize of the tick labels
        plt.rc('ytick', labelsize=SMALL_SIZE)  # fontsize of the tick labels
        plt.rc('legend', fontsize=SMALL_SIZE)  # legend fontsize
        plt.rc('figure', titlesize=BIGGER_SIZE)  # fontsize of the figure title
        lata = self.gencathist(catframe, ax1, ax2, 'all',self.kotybool)
        wx.CallAfter(self.postep.SetValue, 1)

        lata = self.gencathist(catframe, ax3, ax4, 'f', self.kotybool)
        wx.CallAfter(self.postep.SetValue, 2)

        lata = self.gencathist(catframe, ax5, ax6, 'm',self.kotybool)
        wx.CallAfter(self.postep.SetValue, 3)


        fig.savefig(pathname)
        # fig.show()

        self.openfile(pathname)
    def openfile(self, file):
        if platform.system() == 'Darwin':  # macOS
            subprocess.call(('open',file))
        elif platform.system() == 'Windows':  # Windows
            os.startfile(file)
        else:  # linux variants
            subprocess.call(('xdg-open',file))


    def combochange(self,event):
        if self.backup.Value!=self.nadzis:
            self.word1.Enabled=True
            self.histogram1.Enabled=True
            self.genxls.Enabled=True
            self.statcat.Enabled=True
            self.histnew.Enabled = True
            self.histold.Enabled = True
            self.wordNew.Enabled = True
            self.wordOld.Enabled = True
        elif self.newgen[0] and self.kotybool:
            self.word1.Enabled = True
            self.histogram1.Enabled = True
            self.genxls.Enabled = True
            self.statcat.Enabled = True
            self.histnew.Enabled = True
            self.histold.Enabled = True
            self.wordNew.Enabled = True
            self.wordOld.Enabled = True
        elif self.newgen[1] and not self.kotybool:
            self.word1.Enabled=True
            self.histogram1.Enabled=True
            self.genxls.Enabled=True
            self.statcat.Enabled=True
            self.histnew.Enabled = True
            self.histold.Enabled = True
            self.wordNew.Enabled = True
            self.wordOld.Enabled = True
        else:
            self.word1.Enabled=False
            self.histogram1.Enabled=False
            self.genxls.Enabled=False
            self.statcat.Enabled=False
            self.histnew.Enabled=False
            self.histold.Enabled=False
            self.wordNew.Enabled=False
            self.wordOld.Enabled=False



    def switchkotpies(self, event):
        self.listbackups = [ self.nadzis]

        self.backup.Clear()
        self.backup2.Clear()
        # self.backup2.Append( self.nadzis)
        self.backup.Append( self.nadzis)
        self.backup.SetValue( self.nadzis)
        if self.kotybool:
            self.koty.SetValue(False)
            self.psy.SetValue(True)
            self.kotybool = False

            lista = glob.glob(self.my_datadir2 + 'Hau*.kotolizer')
            for k in lista:
                self.listbackups.append(k.split('Hau')[1].split('.kotolizer')[0])
                self.backup.Append(k.split('Hau')[1].split('.kotolizer')[0])
                self.backup2.Append(k.split('Hau')[1].split('.kotolizer')[0])


        else:
            self.psy.SetValue(False)
            self.koty.SetValue(True)
            self.kotybool = True
            lista = glob.glob(self.my_datadir2 + 'Miau*.kotolizer')
            for k in lista:
                self.listbackups.append(k.split('Miau')[1].split('.kotolizer')[0])
                self.backup.Append(k.split('Miau')[1].split('.kotolizer')[0])
                self.backup2.Append(k.split('Miau')[1].split('.kotolizer')[0])

        self.backup2.SetValue(self.listbackups[-1])
        self.combochange(True)

    def parsepaluch(self, event):

        first2 = Thread(target=self.parsepaluch2)
        first2.setDaemon(True)
        first2.start()
        self.nowedane.Enabled = False

    def parsepaluch2(self):
        self.psy.Enabled=False
        self.koty.Enabled=False
        linkMain = 'https://napaluchu.waw.pl/zwierzeta/zwierzeta-do-adopcji/?pet_page='
        if self.kotybool:

            linkEnd = '&pet_species=2'  # Koty
        else:
            linkEnd = '&pet_species=1'  # PSY
        wx.CallAfter(self.postep.SetValue, 10)

        links = self.get_all_cat_links(linkMain, linkEnd)

        linkMain = 'https://napaluchu.waw.pl/zwierzeta/ostatnio-znalezione/?pet_page='
        links = np.concatenate((links, self.get_all_cat_links(linkMain, linkEnd)))
        wx.CallAfter(self.postep.SetRange, len(links) + 1)
        wx.CallAfter(self.postep.SetValue, 1)
        self.catframe = pd.DataFrame(columns=['Nr', 'ID', 'Imię', 'Wiek', 'Przyjęcie', 'Foto', 'Ogłoszenie'])
        opo = 1
        for countlinks, k in enumerate(links):
            CatID, CatName, CatAge, CatIn, CatDesc, CatStat, CatImage = self.get_cat_details(k)
            if CatID != 0:
                new_row = {'Nr': opo, 'ID': CatID, 'Imię': CatName, 'Wiek': CatAge, 'Przyjęcie': CatIn,
                           'Foto': CatImage, 'Ogłoszenie': CatDesc}
                self.catframe = self.catframe.append(new_row, ignore_index=True)
                wx.CallAfter(self.postep.SetValue, countlinks)
                opo += 1
        if self.kotybool:
            self.catframe.to_pickle(self.my_datadir2 + 'Miau' +self.today +'.kotolizer', compression='bz2')
            self.newgen[0]=True
        else:
            self.catframe.to_pickle(self.my_datadir2 + 'Hau' +self.today +'.kotolizer', compression='bz2')
            self.newgen[1]=True

        wx.CallAfter(self.postep.SetValue, len(links) + 1)

        self.word1.Enabled = True
        self.histogram1.Enabled = True
        self.genxls.Enabled = True
        self.statcat.Enabled = True
        self.histnew.Enabled=True
        self.histold.Enabled=True
        self.wordOld.Enabled=True
        self.wordNew.Enabled=True

        # self.nowedane.Enabled=True

        # links2=get_all_cat_links(linkMain,linkEnd)

    def get_cat_details(self, link):
        # ll=links[1]

        soup = self.getsoup('https://napaluchu.waw.pl/' + link)
        CatStat = str(soup.body.find_all('strong')[5])
        #     CatStat+='\n\n'
        #     CatStat+=str(soup.body.find_all('strong')[8])
        CatID = int(str(soup.body.find_all('h2')[0]).split(' ')[-1].split('>')[1].split('/')[0])
        CatName = str(soup.body.find_all('h2')[0]).split(' ')[0].split('>')[1] + '\n' + \
                  str(soup.body.find_all('strong')[2]).split('>')[1].split('<')[0]
        CatAge = str(str(soup.body.find_all('strong')[1]).split('>')[1].split('</')[0])
        # soup.body.find('div', {'class':"autocontainer"})
        CatIn = (str(soup.body.find_all('strong')[6]).split('>')[1].split('<')[0]) + '\n' + \
                str(soup.body.find_all('strong')[7]).split('>')[1].split('<')[0]
        CatIn += '\n\nZdjęć:' + str(len(soup.body.find_all('div', class_='pet-detail-gallery-column-photo')))
        if 'facebook' not in str(soup.body.find_all('strong')[8]):
            CatIn += '\n\n' + str(soup.body.find_all('strong')[8]).replace('<strong>', '').replace('</strong>', '')

        #     soup.body.find_all('strong')[7]

        ImagePath = requests.get(
            str(soup.body.find_all('img', class_='pet-detail-main-image')).split('data-src=')[1].split(' ')[0].replace(
                '"', ''))
        fixed_height = 150
        image = Image.open(BytesIO(ImagePath.content))
        height_percent = (fixed_height / float(image.size[0]))
        width_size = int((float(image.size[1]) * float(height_percent)))
        CatImage = image.resize((fixed_height, width_size), PIL.Image.NEAREST)
        CatDesc = str(soup.body.find_all('div', class_='pet-description')).split('>')[1].split('<')[0].split(
            'Zapraszamy do zapoznania się z ankietą')[0]
        if "pilne" in CatStat:
            CatIn += '\n\nPILNE'
        if 'Ochota na Kota' in CatStat:
            CatIn += '\n\nOchota na Kota'

        if not ("tymczaso" or "fundacja") in CatStat:
            return CatID, CatName, CatAge, CatIn, CatDesc, CatStat, CatImage
        else:
            return 0, 0, 0, 0, 0, CatStat, 0

    def get_links_to_cats(self, soup, links):
        # sciaga linki z danej podstrony
        # links = []
        for link in soup.findAll('a', attrs={'href': re.compile("/pet/")}):
            links.append(link.get('href'))
        return links

    def get_max_page(self, soup):
        # podaje liczbe podstron do sciagania
        nums = []
        for link in soup.findAll('a', attrs={'href': re.compile("pet_page=")}):
            nums.append(int(link.get('href').split('?')[1].split('=')[1].split('&')[0]))
        return max(nums)

    def getsoup(self, link):
        odpowiedz = requests.get(link)
        return BeautifulSoup(odpowiedz.text, 'html.parser')

    def get_all_cat_links(self, linkMain, linkEnd):
        odpowiedz = requests.get(linkMain + str(1) + linkEnd)
        soup = BeautifulSoup(odpowiedz.text, 'lxml')
        maxcount = self.get_max_page(soup)
        links = []
        for k in range(1, maxcount + 1):
            soup = self.getsoup(linkMain + str(k) + linkEnd)
            links = self.get_links_to_cats(soup, links)
        return np.unique(links)

    def getframecur(self, b):
        if b.Value != self.nadzis:
            if self.kotybool:
                df = pd.read_pickle(self.my_datadir2 + 'Miau' + b.Value + '.kotolizer',
                                    compression='bz2')
            else:
                df = pd.read_pickle(self.my_datadir2 + 'Hau' +  b.Value + '.kotolizer',
                                    compression='bz2')
            catframe = df.sort_values('ID')

        else:
            catframe = self.catframe.sort_values('ID')
        return catframe

    def comparedf(self,one, two):
        kiki = pd.concat([one, two]).drop_duplicates(subset='ID', keep=False)
        new = pd.merge(one, kiki, how='inner', on='ID')
        df3 = new[['Nr_x', 'ID', 'Imię_x', 'Wiek_x', 'Przyjęcie_x', 'Foto_x', 'Ogłoszenie_x']].rename(
            columns={'Przyjęcie_x': 'Przyjęcie',
                     'Ogłoszenie_x': 'Ogłoszenie',
                     'Nr_x': 'Nr', 'Imię_x': 'Imię',
                     'Wiek_x': 'Wiek', 'Foto_x': 'Foto'})
        return df3
    def genword(self, event):

        with wx.FileDialog(self, "Save file", wildcard="DOCX files (*.docx)|*.docx",
                           style=wx.FD_SAVE | wx.FD_OVERWRITE_PROMPT) as fileDialog:
            if fileDialog.ShowModal() == wx.ID_CANCEL:
                return  # the user changed their mind

            # save the current contents in the file
            pathname = fileDialog.GetPath()
            catframe=self.getframecur(self.backup)

            first2 = Thread(target=self.createwordfile, args=(pathname,catframe,))
            first2.setDaemon(True)
            first2.start()
            print(pathname)
            # self.createwordfile(self.catframe, pathname)
        print('Done')
    def genwordNew(self, event):
        catframe = self.getframecur(self.backup)
        catframe2 = self.getframecur(self.backup2)
        Nowekoty = self.comparedf(catframe, catframe2)
        if Nowekoty.shape[0] == 0:
            dlg = wx.MessageDialog(self, 'Nie przybyło zwierzaków', 'Nie ma nowych zwierząt w porówaniu').ShowModal()
        else:
            with wx.FileDialog(self, "Save file", wildcard="DOCX files (*.docx)|*.docx",
                               style=wx.FD_SAVE | wx.FD_OVERWRITE_PROMPT) as fileDialog:
                if fileDialog.ShowModal() == wx.ID_CANCEL:
                    return  # the user changed their mind

                # save the current contents in the file
                pathname = fileDialog.GetPath()
                # catframe=self.getframecur(self.backup)
                # catframe2=self.getframecur(self.backup2)
                # Nowekoty = self.comparedf(catframe, catframe2)


                first2 = Thread(target=self.createwordfile, args=(pathname,Nowekoty,))
                first2.setDaemon(True)
                first2.start()
                print(pathname)
            # self.createwordfile(self.catframe, pathname)
        print('Done')
    def genwordOld(self, event):
        catframe = self.getframecur(self.backup)
        catframe2 = self.getframecur(self.backup2)
        # Nowekoty = comparedf(catframe, zeszly_tydzien)
        Zniknietekoty = self.comparedf(catframe2, catframe)
        if Zniknietekoty.shape[0] == 0:
            dlg = wx.MessageDialog(self, 'Nie ubyło zwierzaków', 'Stan zwierząt nie zmniejszył się').ShowModal()
        else:
            with wx.FileDialog(self, "Save file", wildcard="DOCX files (*.docx)|*.docx",
                               style=wx.FD_SAVE | wx.FD_OVERWRITE_PROMPT) as fileDialog:
                if fileDialog.ShowModal() == wx.ID_CANCEL:
                    return  # the user changed their mind

                # save the current contents in the file
                pathname = fileDialog.GetPath()

                first2 = Thread(target=self.createwordfile, args=(pathname,Zniknietekoty,))
                first2.setDaemon(True)
                first2.start()
                print(pathname)
            # self.createwordfile(self.catframe, pathname)
        print('Done')

    def createwordfile(self, pathname,catframe):

        try:
            os.mkdir(self.my_datadir2+'thumb')
        except:
            pass
        # open an existing document
        doc = docx.Document()

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(8)

        section = doc.sections[0]
        section.page_height = Mm(210)
        section.page_width = Mm(297)
        t = doc.add_table(catframe.shape[0] + 1, catframe.shape[1])

        # add the header rows.
        # wx.CallAfter(self.postep.SetRange, 1)

        # pub.sendMessage('postep1', msg=(catframe.shape[0]))
        # pub.sendMessage('postep2', msg=0)
        wx.CallAfter(self.postep.SetRange, (catframe.shape[0]))

        for j in range(catframe.shape[-1]):
            t.cell(0, j).text = catframe.columns[j]

        # add the rest of the data frame
        for i in range(catframe.shape[0]):
            # pub.sendMessage('postep2', msg=i+1)
            wx.CallAfter(self.postep.SetValue, i + 1)

            for j in range(catframe.shape[-1]):

                if j == 5:
                    catframe.values[i, j].save(self.my_datadir2+'thumb/' + str(i) + '.jpg', "JPEG")
                    paragraph = t.cell(i + 1, j).paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(self.my_datadir2+'thumb/' + str(i) + '.jpg')
                elif j == 0:
                    t.cell(i + 1, j).text = str(i + 1)


                else:
                    t.cell(i + 1, j).text = str(catframe.values[i, j])

                if 'PILNE' in catframe.values[i, 4]:
                    t.cell(i + 1, j)._tc.get_or_add_tcPr().append(
                        parse_xml(r'<w:shd {} w:fill="f8ddd7"/>'.format(nsdecls('w'))))
            if 'Ochota na Kota' in catframe.values[i, 4]:
                t.cell(i + 1, 2)._tc.get_or_add_tcPr().append(
                    parse_xml(r'<w:shd {} w:fill="f0d7f8"/>'.format(nsdecls('w'))))

        ####f0d7f8
        t.allow_autofit = True
        t.columns[0].width = Cm(0.7)
        t.columns[1].width = Cm(1.5)
        t.columns[2].width = Cm(1.5)
        t.columns[3].width = Cm(1.2)
        t.columns[4].width = Cm(2)
        t.columns[5].width = Cm(6)
        t.columns[6].width = Cm(10)
        doc.save(pathname)
        shutil.rmtree(self.my_datadir2+'thumb/')
        self.openfile(pathname)


    def get_datadir(self) -> pathlib.Path:

        """
        Returns a parent directory path
        where persistent application data can be stored.

        # linux: ~/.local/share
        # macOS: ~/Library/Application Support
        # windows: C:/Users/<USER>/AppData/Roaming
        """

        home = pathlib.Path.home()

        if sys.platform == "win32":
            return home / "AppData/Roaming"
        elif sys.platform == "linux":
            return home / ".local/share"
        elif sys.platform == "darwin":
            return home / "Library/Application Support"

    # create your program's directory


    def findbackups(self):
        # Dir=self.path
        self.listbackups = []
        self.listbackups.append('Current')
        self.backup.Clear()
        self.backup.Append('Current')
        self.backup.SetValue('Current')
        a = 0
        for name in os.listdir(self.path):
            if (('Reconstructed_back' in name)) and os.path.isfile(
                    self.path + '/' + name + '/params.nmr'):
                self.listbackups.append(name.replace('Reconstructed_back_at', ''))
                self.backup.Append(name.replace('Reconstructed_back_at', ''))
                a += 1
        if a > 0:
            self.backup.Enable()
            self.backuplistNOTgenerated = False


class CatFrame(wx.Frame):
    def __init__(self):
        wx.Frame.__init__(self, None, title="Kotolizator 3000 - Analizator stanu zwierząt w schronisku na Paluchu",
                          size=(800, 500))
        nolog = wx.LogNull()  # windows alert fix

        if platform.system() == 'Linux' and getattr(sys, 'frozen', False):
            IconPath = sys._MEIPASS + '/icon.ico'
            if float(wx.__version__[0]) < 4:
                self.SetIcon(wx.IconFromBitmap(wx.Bitmap(IconPath, wx.BITMAP_TYPE_ANY)))
            else:
                self.SetIcon(wx.Icon(wx.Bitmap(IconPath, wx.BITMAP_TYPE_ANY)))
        else:
            try:  # Trend.Resources is for code used from whl distribution
                IconPath = os.path.abspath(resource_filename('analizatorkotow', 'icon.ico'))
                print(IconPath)
                if float(wx.__version__[0]) < 4:
                    self.SetIcon(wx.IconFromBitmap(wx.Bitmap(IconPath, wx.BITMAP_TYPE_ANY)))
                else:
                    self.SetIcon(wx.Icon(wx.Bitmap(IconPath, wx.BITMAP_TYPE_ANY)))
            except:  # If code is used from source code
                IconPath = 'icon.ico'
                if float(wx.__version__[0]) < 4:
                    self.SetIcon(wx.IconFromBitmap(wx.Bitmap(IconPath, wx.BITMAP_TYPE_ANY)))
                else:
                    self.SetIcon(wx.Icon(wx.Bitmap(IconPath, wx.BITMAP_TYPE_ANY)))
        del nolog
        self.p = wx.Panel(self)  # creates panel
        self.nb = GeneralLayout(self.p)
        self.verSizer = wx.BoxSizer(wx.VERTICAL)
        self.verSizer.Add(self.nb, 1, wx.ALL | wx.EXPAND, 0)
        self.firstimeclick = True
        self.statusbar = self.CreateStatusBar()
        self.Centre()
        self.p.SetSizer(self.verSizer)
        self.p.Fit()
        self.statusbar.SetStatusText(
            'Mateusz Urbańczyk z dedykacją dla Karoliny')

    def updatecontrolers(self, msg):
        e = wx.SizeEvent(self.GetSize())
        self.ProcessEvent(e)
        self.Fit()
